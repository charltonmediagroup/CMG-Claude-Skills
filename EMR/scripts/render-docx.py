"""Render a single-publication Earned Media Report DOCX from cache/<slug>-content.json + curated dataset.

Usage:
  python EMR/scripts/render-docx.py --site SBR --slug dbs

This script handles SINGLE-PUBLICATION reports only. Combined-mode reports
(2+ publications) are rendered by the skill at runtime via anthropic-skills:docx,
not by this script. See SKILL.md for the combined-mode workflow.

Layout matches the SKILL.md HTML template:
  - Title block (Heading 1) + meta (Prepared by / Period / Date)
  - Executive Summary (Heading 2 + 2 paragraphs)
  - Financial Summary table
  - By Year table
  - By Category table
  - Strategic Observations (Heading 2 + bulleted list)
  - Article Inventory (Heading 2 + intro paragraph + N-row table with hyperlinked titles)

Recommendations were removed from every report; this script no longer emits that section.

Typography: body 11pt, headings 14pt. Hyperlinks blue (#1a4ea0).
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


BODY_PT = 11
HEAD_PT = 14
LINK_COLOR = "1a4ea0"


def set_run_font(run, *, size_pt: int = BODY_PT, bold: bool = False, color_hex: str | None = None):
    run.font.name = "Helvetica"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def add_heading(doc: Document, text: str, level: int = 2):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size_pt=HEAD_PT, bold=True)
    return p


def add_para(doc: Document, text: str = "", *, bold: bool = False) -> "Paragraph":
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold)
    return p


def add_runs(p, segments: list[tuple[str, dict]]):
    """Add multiple runs to a paragraph: each segment is (text, style_kwargs)."""
    for text, kw in segments:
        run = p.add_run(text)
        set_run_font(run, size_pt=kw.pop("size_pt", BODY_PT), **kw)


def add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink run into the paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Helvetica")
    rFonts.set(qn("w:hAnsi"), "Helvetica")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(BODY_PT * 2))
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_COLOR)
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def style_table_cell(cell, *, bold: bool = False, header: bool = False, num: bool = False):
    for p in cell.paragraphs:
        for run in p.runs:
            set_run_font(run, bold=bold or header)
        if num:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    if header:
        # Light grey background
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "f2f2f2")
        tcPr.append(shd)


def add_table(doc: Document, header: list[str], rows: list[list], *, num_cols: list[int] | None = None,
              bold_total_row: bool = False):
    num_cols = num_cols or []
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    # header
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_table_cell(cell, header=True, num=(i in num_cols))
    # body
    for r, row in enumerate(rows):
        is_total = bold_total_row and r == len(rows) - 1
        for i, v in enumerate(row):
            cell = table.rows[r + 1].cells[i]
            cell.text = str(v)
            style_table_cell(cell, bold=is_total, num=(i in num_cols))
            if is_total:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "f7f7f7")
                tcPr.append(shd)
    return table


def fmt_dollars(n: int) -> str:
    return f"USD {n:,}"


def render(site: str, slug: str, workspace_root: Path) -> Path:
    out_dir = workspace_root / "EMR" / "output" / site
    cache = out_dir / ".cache"
    content_path = cache / f"{slug}-content.json"
    if not content_path.exists():
        raise FileNotFoundError(content_path)
    content = json.loads(content_path.read_text(encoding="utf-8"))

    doc = Document()
    # Set default font on Normal style
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(BODY_PT)
    # Heading styles
    for hname, hpt in (("Heading 1", HEAD_PT), ("Heading 2", HEAD_PT)):
        s = doc.styles[hname]
        s.font.name = "Helvetica"
        s.font.size = Pt(hpt)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # ---- Title block ----
    add_heading(doc, f"{content['brand']} — Earned Media Report", level=1)
    meta = doc.add_paragraph()
    add_runs(meta, [
        ("Prepared by: ", {"bold": True}),
        (content["publication"], {}),
        ("\n", {}),
        ("Period: ", {"bold": True}),
        (content["period"], {}),
        ("\n", {}),
        ("Date: ", {"bold": True}),
        (content["report_date"], {}),
    ])

    # ---- Executive Summary ----
    add_heading(doc, "Executive Summary")
    # Para 1 with inline bolds
    p1 = doc.add_paragraph()
    add_runs(p1, [
        ("We at ", {}),
        (content["publication"], {}),
        (" have published ", {}),
        (str(content["total_pieces"]), {"bold": True}),
        (" unique editorial pieces in which DBS, POSB, or one of the group's sub-brands is the direct subject. The coverage spans ", {}),
        ("22 January 2010 through 30 April 2026", {"bold": True}),
        (" — sixteen years and three months of continuous editorial attention. At our agreed flat valuation of USD 1,800 per piece, this corpus represents a total Earned Media Value of ", {}),
        (f"USD {content['emv']}", {"bold": True}),
        (".", {}),
    ])
    add_para(doc, content["exec_summary_para_2"])

    # ---- Financial Summary ----
    add_heading(doc, "Financial Summary")
    fin_rows = [
        ["Total unique editorial pieces", str(content["total_pieces"])],
        ["Flat media value per piece", "USD 1,800"],
        ["Total Earned Media Value", f"USD {content['emv']}"],
        ["Coverage window", content["coverage_window"]],
    ]
    t = add_table(doc, ["Metric", "Value"], fin_rows)
    # Bold value column for first three rows
    for r in range(1, 4):
        for run in t.rows[r].cells[1].paragraphs[0].runs:
            run.bold = True

    # ---- By Year ----
    add_heading(doc, "By Year")
    by_year = content["byYear"]
    yrs = sorted(by_year.keys())
    rows = []
    total = 0
    total_emv = 0
    for y in yrs:
        n = by_year[y]
        e = n * 1800
        total += n
        total_emv += e
        rows.append([y, str(n), f"${e:,}"])
    rows.append(["Total", str(total), f"${total_emv:,}"])
    add_table(doc, ["Year", "Pieces", "EMV (USD)"], rows, num_cols=[1, 2], bold_total_row=True)

    # ---- By Category ----
    add_heading(doc, "By Category")
    by_cat = content["byCat"]
    cat_rows = []
    cat_total = sum(by_cat.values())
    for cat, n in sorted(by_cat.items(), key=lambda kv: -kv[1]):
        share = (n / cat_total * 100) if cat_total else 0
        cat_rows.append([cat, str(n), f"{share:.1f}%"])
    add_table(doc, ["Category", "Pieces", "Share"], cat_rows, num_cols=[1, 2])

    # ---- Strategic Observations ----
    add_heading(doc, "Strategic Observations")
    for obs in content["observations"]:
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, [
            (obs["thesis"] + " ", {"bold": True}),
            (obs["body"], {}),
        ])

    # ---- Recommendations ----
    # Recommendations were removed from every report (single-pub and combined)
    # per the multi-publication redesign. The skill no longer composes a
    # "recommendations_framing" or "recommendations" field in content.json,
    # so this section is intentionally empty. Do not re-add it without
    # updating SKILL.md and the content schema first.

    # ---- Article Inventory ----
    add_heading(doc, "Article Inventory")
    add_para(doc,
             f"All {content['total_pieces']} unique editorial pieces published on {content['publication']} "
             f"in which {content['brand']} or one of its sub-brands is the direct subject. Listed most recent first. "
             f"Each carries the agreed flat valuation of USD 1,800.")

    # Build inventory table (header + N rows)
    arts = content["articles"]
    inv = doc.add_table(rows=1 + len(arts), cols=4)
    inv.style = "Table Grid"
    headers = ["#", "Date", "Article", "Category"]
    for i, h in enumerate(headers):
        cell = inv.rows[0].cells[i]
        cell.text = h
        style_table_cell(cell, header=True, num=(i in (0,)))
    for idx, a in enumerate(arts, start=1):
        row = inv.rows[idx].cells
        row[0].text = str(idx)
        row[1].text = a.get("date") or ""
        # Article cell: hyperlinked title
        art_cell = row[2]
        art_cell.text = ""  # clear default empty paragraph text
        para = art_cell.paragraphs[0]
        if a.get("url"):
            add_hyperlink(para, a["url"], a.get("title") or a["url"])
        else:
            run = para.add_run(a.get("title") or "")
            set_run_font(run)
        row[3].text = a.get("category") or ""
        for i, cell in enumerate(row):
            style_table_cell(cell, num=(i == 0))

    # ---- Save ----
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{slug}.docx"
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--site", required=True)
    ap.add_argument("--slug", required=True)
    ap.add_argument("--workspace-root", default=".")
    args = ap.parse_args()
    out = render(args.site, args.slug, Path(args.workspace_root))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
